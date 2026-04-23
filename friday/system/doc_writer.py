"""
Document Writer Module
Creates Word documents with extracted equations and solutions.
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("python-docx is required. Install it with: pip install python-docx")


logger = logging.getLogger(__name__)


class DocumentWriter:
    """Create and write Word documents with equation solutions."""

    def __init__(self, log_level: int = logging.INFO):
        """
        Initialize DocumentWriter.
        
        Args:
            log_level: Logging level for the writer
        """
        self.logger = logging.getLogger(self.__class__.__name__)
        self.logger.setLevel(log_level)

    def create_solution_document(
        self,
        output_path: str,
        equations: List[str],
        solutions: List[Dict[str, Any]],
        pdf_source: Optional[str] = None,
        title: str = "Equation Solutions"
    ) -> bool:
        """
        Create a Word document with equations and solutions.
        
        Args:
            output_path: Path to save the document
            equations: List of extracted equations
            solutions: List of solution dictionaries
            pdf_source: Source PDF filename
            title: Document title
            
        Returns:
            True if document created successfully
        """
        try:
            self.logger.info(f"Creating document: {output_path}")
            
            # Create document
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            self._add_metadata(doc, pdf_source)
            
            # Add equations and solutions
            if equations and solutions:
                self._add_equations_and_solutions(doc, equations, solutions)
            else:
                doc.add_paragraph("No equations found to solve.")
            
            # Save document
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            doc.save(str(output_file))
            self.logger.info(f"Document saved successfully: {output_path}")
            return True

        except Exception as e:
            error_msg = f"Error creating document: {str(e)}"
            self.logger.error(error_msg)
            return False

    def create_system_solution_document(
        self,
        output_path: str,
        system_equations: List[str],
        system_solutions: List[Dict[str, str]],
        pdf_source: Optional[str] = None
    ) -> bool:
        """
        Create a Word document for system of equations solutions.
        
        Args:
            output_path: Path to save the document
            system_equations: List of equations in the system
            system_solutions: List of solution dictionaries
            pdf_source: Source PDF filename
            
        Returns:
            True if document created successfully
        """
        try:
            self.logger.info(f"Creating system solution document: {output_path}")
            
            doc = Document()
            
            # Add title
            title_para = doc.add_heading("System of Equations - Solutions", level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            self._add_metadata(doc, pdf_source)
            
            # Add equations section
            doc.add_heading("Equations:", level=2)
            for i, eq in enumerate(system_equations, 1):
                doc.add_paragraph(f"{i}. {eq}", style='List Number')
            
            # Add solutions section
            doc.add_heading("Solution:", level=2)
            
            if system_solutions:
                for i, sol in enumerate(system_solutions, 1):
                    if isinstance(sol, dict):
                        doc.add_paragraph(f"Solution {i}:", style='Heading 3')
                        for var, val in sol.items():
                            para = doc.add_paragraph(f"  {var} = {val}")
                            self._format_solution_paragraph(para)
                    else:
                        doc.add_paragraph(f"Solution {i}: {str(sol)}")
            else:
                doc.add_paragraph("No solutions found for this system of equations.")
            
            # Save document
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            doc.save(str(output_file))
            self.logger.info(f"System solution document saved: {output_path}")
            return True

        except Exception as e:
            error_msg = f"Error creating system solution document: {str(e)}"
            self.logger.error(error_msg)
            return False

    def _add_metadata(self, doc: Document, pdf_source: Optional[str] = None):
        """Add metadata section to document."""
        metadata_para = doc.add_paragraph()
        
        # Add timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        run = metadata_para.add_run(f"Generated: {timestamp}\n")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add PDF source if provided
        if pdf_source:
            run = metadata_para.add_run(f"Source: {pdf_source}\n")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add separator
        doc.add_paragraph()

    def _add_equations_and_solutions(
        self,
        doc: Document,
        equations: List[str],
        solutions: List[Dict[str, Any]]
    ):
        """Add equations and their solutions to document."""
        doc.add_heading("Equations and Solutions:", level=2)
        
        for i, (eq, sol) in enumerate(zip(equations, solutions), 1):
            # Add equation number and original equation
            doc.add_heading(f"Equation {i}", level=3)
            
            eq_para = doc.add_paragraph()
            run = eq_para.add_run(f"Original: {eq}")
            run.font.bold = True
            
            # Add solution status
            status = sol.get("status", "unknown")
            
            if status == "success":
                solutions_list = sol.get("solutions", [])
                
                if solutions_list:
                    doc.add_paragraph("Solution(s):", style='List Bullet')
                    
                    for j, solution in enumerate(solutions_list, 1):
                        if isinstance(solution, dict):
                            for var, val in solution.items():
                                sol_para = doc.add_paragraph(
                                    f"{var} = {val}",
                                    style='List Bullet 2'
                                )
                                self._format_solution_paragraph(sol_para)
                        else:
                            doc.add_paragraph(str(solution), style='List Bullet 2')
                else:
                    doc.add_paragraph(
                        "No solutions found (equation may be inconsistent or trivial).",
                        style='List Bullet'
                    )
            
            else:
                error_msg = sol.get("message", "Unknown error")
                error_para = doc.add_paragraph(f"Error: {error_msg}")
                error_run = error_para.runs[0] if error_para.runs else None
                if error_run:
                    error_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Add separator
            doc.add_paragraph()

    def _format_solution_paragraph(self, paragraph):
        """Apply formatting to solution paragraph."""
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(11)

    def append_to_document(
        self,
        doc_path: str,
        equations: List[str],
        solutions: List[Dict[str, Any]],
        section_title: str = "New Equations"
    ) -> bool:
        """
        Append equations and solutions to existing document.
        
        Args:
            doc_path: Path to existing document
            equations: List of equations to append
            solutions: List of solutions
            section_title: Title for new section
            
        Returns:
            True if append successful
        """
        try:
            doc_file = Path(doc_path)
            
            if not doc_file.exists():
                self.logger.warning(f"Document not found: {doc_path}, creating new")
                return self.create_solution_document(doc_path, equations, solutions, title=section_title)
            
            # Open existing document
            doc = Document(str(doc_file))
            
            # Add new section
            doc.add_page_break()
            doc.add_heading(section_title, level=2)
            
            # Add equations and solutions
            self._add_equations_and_solutions(doc, equations, solutions)
            
            # Save
            doc.save(str(doc_file))
            self.logger.info(f"Appended to document: {doc_path}")
            return True

        except Exception as e:
            error_msg = f"Error appending to document: {str(e)}"
            self.logger.error(error_msg)
            return False

    def create_summary_document(
        self,
        output_path: str,
        summary_data: Dict[str, Any],
        pdf_source: Optional[str] = None
    ) -> bool:
        """
        Create a comprehensive summary document.
        
        Args:
            output_path: Path to save document
            summary_data: Dictionary with 'equations_extracted', 'equations_solved', 'solutions'
            pdf_source: Source PDF filename
            
        Returns:
            True if document created successfully
        """
        try:
            doc = Document()
            
            # Title
            title_para = doc.add_heading("PDF Equation Solver - Summary Report", level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            self._add_metadata(doc, pdf_source)
            
            # Summary statistics
            doc.add_heading("Summary:", level=2)
            
            summary_para = doc.add_paragraph()
            summary_para.add_run(
                f"Equations Extracted: {summary_data.get('equations_extracted', 0)}\n"
            )
            summary_para.add_run(
                f"Equations Solved: {summary_data.get('equations_solved', 0)}\n"
            )
            
            # Processing time if available
            if 'processing_time' in summary_data:
                summary_para.add_run(f"Processing Time: {summary_data['processing_time']:.2f}s\n")
            
            doc.add_paragraph()
            
            # Detailed solutions
            if 'solutions' in summary_data:
                doc.add_heading("Solutions:", level=2)
                self._add_equations_and_solutions(
                    doc,
                    summary_data.get('equations', []),
                    summary_data['solutions']
                )
            
            # Save
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_file))
            
            self.logger.info(f"Summary document created: {output_path}")
            return True

        except Exception as e:
            error_msg = f"Error creating summary document: {str(e)}"
            self.logger.error(error_msg)
            return False
